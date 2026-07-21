import json
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- Colores de Marca Dra. Fiona Xacur ---
HEX_PRIMARY = "0F4C81"   # Azul Océano / Oftalmológico
HEX_ACCENT = "D59704"    # Dorado Láser
HEX_TEXT = "1E293B"      # Gris Oscuro Elegante
HEX_LIGHT_BG = "F0F7FA"  # Azul muy claro para fondos de tablas y callouts
HEX_WHITE = "FFFFFF"
HEX_BORDER = "CBD5E1"    # Borde suave

COLOR_PRIMARY = RGBColor(15, 76, 129)
COLOR_ACCENT = RGBColor(213, 151, 4)
COLOR_TEXT = RGBColor(30, 41, 59)
COLOR_WHITE = RGBColor(255, 255, 255)

# --- Helper XML Functions for python-docx Styling ---
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180): # in dxa (1 pt = 20 dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>''')
    tcPr.append(tcMar)

def set_callout_borders(cell, border_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="none"/>
        <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
    </w:tcBorders>''')
    tcPr.append(tcBorders)

def set_table_borders(table, color=HEX_BORDER):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
        <w:left w:val="none"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>''')
    tblPr.append(borders)

def add_formatted_text(paragraph, text, default_color=COLOR_TEXT, italic=False):
    if not text:
        return
    parts = str(text).split("**")
    for idx, part in enumerate(parts):
        if not part and idx == 0:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = default_color
        run.italic = italic
        if idx % 2 == 1:
            run.bold = True

def add_paragraph_with_spacing(doc, text="", style='Normal', space_after=6, space_before=0, bullet=False):
    style_name = 'List Bullet' if bullet else style
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if text:
        add_formatted_text(p, text)
    return p

def add_custom_heading(doc, text, level, space_before=14, space_after=6, keep_with_next=True):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = keep_with_next
    
    run = h.add_run(text)
    run.font.name = 'Calibri Light' if level == 1 else 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_PRIMARY
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_ACCENT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_PRIMARY
    return h

def add_callout(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, HEX_LIGHT_BG)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    set_callout_borders(cell, HEX_ACCENT)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        run_title = p.add_run(title + "\n")
        run_title.font.name = 'Calibri'
        run_title.font.size = Pt(11)
        run_title.bold = True
        run_title.font.color.rgb = COLOR_ACCENT
        
    add_formatted_text(p, text, default_color=COLOR_PRIMARY, italic=True)
    add_paragraph_with_spacing(doc, space_after=12)

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    json_path = os.path.join(script_dir, "../../data.json")
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    doctor = data["doctor"]
    services = data["services"]
    symptoms = data["symptoms"]
    diseases = data["diseases"]
    privacy = data["privacyPolicy"]
    
    # ----------------------------------------------------
    # 1. GENERACIÓN DEL DOCUMENTO WORD (.docx)
    # ----------------------------------------------------
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run(f"{doctor['title']} {doctor['name']}  |  Guía de Contenido Web")
        f_run.font.name = 'Calibri'
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(120, 120, 120)

    # 1.1 Portada Elegante
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(16)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t_run = title_p.add_run("GUÍA COMPLETA DE CONTENIDO WEB")
    t_run.font.name = 'Calibri Light'
    t_run.font.size = Pt(26)
    t_run.bold = True
    t_run.font.color.rgb = COLOR_PRIMARY
    
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_after = Pt(16)
    l_run = line_p.add_run("―" * 25)
    l_run.font.color.rgb = COLOR_ACCENT
    l_run.bold = True
    
    doc_title_p = doc.add_paragraph()
    doc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title_p.paragraph_format.space_after = Pt(100)
    dt_run = doc_title_p.add_run("Estructura Consolidada de Textos, Servicios, Síntomas y Enfermedades Oculares")
    dt_run.font.name = 'Calibri'
    dt_run.font.size = Pt(13)
    dt_run.font.color.rgb = COLOR_TEXT
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(6)
    dr_name_run = info_p.add_run(f"{doctor['title']} {doctor['name']}\n")
    dr_name_run.font.name = 'Calibri'
    dr_name_run.font.size = Pt(16)
    dr_name_run.bold = True
    dr_name_run.font.color.rgb = COLOR_PRIMARY
    
    dr_title_run = info_p.add_run(f"{doctor['specialty']}\n")
    dr_title_run.font.name = 'Calibri'
    dr_title_run.font.size = Pt(11)
    dr_title_run.font.color.rgb = COLOR_TEXT
    
    dr_sub_run = info_p.add_run(f"{doctor['subspecialty']}")
    dr_sub_run.font.name = 'Calibri'
    dr_sub_run.font.size = Pt(10)
    dr_sub_run.font.color.rgb = COLOR_TEXT
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(40)
    date_run = date_p.add_run(f"{doctor['city']}, {doctor['state']}  |  {datetime.date.today().strftime('%d de %B de %Y')}")
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(125, 125, 125)
    
    doc.add_page_break()

    # 1.2 Sección 1: Perfil de la Doctora
    add_custom_heading(doc, "1. Perfil Profesional de la Doctora", level=1, space_before=18, space_after=12)
    add_paragraph_with_spacing(doc, doctor["bio"])
    
    add_custom_heading(doc, "Filosofía Médica", level=2, space_before=12, space_after=6)
    add_callout(doc, f'"{doctor["philosophy"]}"', title="Misión & Filosofía")
    
    add_custom_heading(doc, "Datos del Consultorio e Información General", level=2, space_before=12, space_after=6)
    
    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_table)
    
    general_details = [
      ("Dirección de Consultorios", doctor["address"]),
      ("Ciudad y Estado", f"{doctor['city']}, {doctor['state']}"),
      ("Teléfono de Contacto", doctor["phone"]),
      ("WhatsApp de Citas", doctor["whatsapp"]),
      ("Correo Electrónico", doctor["email"]),
      ("Costo de Consulta", f"${doctor['consultationPrice']} MXN"),
      ("Métodos de Pago", ", ".join(doctor["paymentMethods"])),
      ("Aseguradoras / Convenios", ", ".join(doctor["insurances"])),
      ("Horarios de Atención", doctor["schedule"]),
      ("Cédula Profesional", doctor["cedula"]),
      ("Cédula de Especialidad", doctor["cedulaEspecialidad"]),
      ("Registro COFEPRIS", doctor["cofepris"]),
    ]
    
    for label, val in general_details:
        row = info_table.add_row()
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        set_cell_background(cell_lbl, HEX_LIGHT_BG)
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(0)
        p_lbl.paragraph_format.line_spacing = 1.15
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Calibri'
        run_lbl.bold = True
        run_lbl.font.color.rgb = COLOR_PRIMARY
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(0)
        p_val.paragraph_format.line_spacing = 1.15
        run_val = p_val.add_run(val)
        run_val.font.name = 'Calibri'
        run_val.font.color.rgb = COLOR_TEXT
        
    add_paragraph_with_spacing(doc, space_after=12)
    
    add_custom_heading(doc, "Educación y Formación Académica", level=2, space_before=12, space_after=6)
    for edu in doctor["education"]:
        p = add_paragraph_with_spacing(doc, bullet=True, space_after=4)
        p.add_run(f"**{edu['degree']}** — {edu['institution']}")
        if edu.get("year"):
            p.add_run(f" ({edu['year']})")
            
    add_custom_heading(doc, "Certificaciones y Colegios Médicos", level=2, space_before=12, space_after=6)
    for cert in doctor["certifications"]:
        p = add_paragraph_with_spacing(doc, bullet=True, space_after=4)
        p.add_run(f"**{cert['name']}** — {cert['institution']}")
            
    add_custom_heading(doc, "Trayectoria y Experiencia", level=2, space_before=12, space_after=6)
    for exp in doctor["experience"]:
        p = add_paragraph_with_spacing(doc, space_after=6)
        run_exp = p.add_run(f"• **{exp['year']} - {exp['title']}:** ")
        run_exp.font.name = 'Calibri'
        run_exp.font.color.rgb = COLOR_TEXT
        p.add_run(exp["description"])

    doc.add_page_break()

    # 1.3 Sección 2: Servicios Médicos
    add_custom_heading(doc, "2. Portafolio de Servicios y Procedimientos Médicos", level=1, space_before=18, space_after=12)
    
    for idx, s in enumerate(services):
        add_custom_heading(doc, f"Servicio {idx+1}: {s['name']}", level=2, space_before=14, space_after=8)
        add_paragraph_with_spacing(doc, f"**Descripción corta:** {s['description']}")
        add_paragraph_with_spacing(doc, f"**Descripción detallada:** {s['longDescription']}")
        
        add_custom_heading(doc, "Ficha Técnica del Procedimiento", level=3, space_before=10, space_after=4)
        spec_table = doc.add_table(rows=0, cols=2)
        spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(spec_table)
        
        specs = [
          ("Tipo de Servicio", s["type"].upper()),
          ("Tipo de Anestesia", s["anesthesiaType"]),
          ("Duración Estimada", s["duration"]),
          ("Tiempo de Recuperación", s["recoveryTime"]),
          ("¿Es doloroso?", "Sí" if s["isPainful"] else "No"),
          ("Rango de Precios", s["priceRange"]),
        ]
        for k, v in s["technicalSpecs"].items():
            specs.append((k, str(v)))
            
        for label, val in specs:
            row = spec_table.add_row()
            cell_lbl, cell_val = row.cells[0], row.cells[1]
            cell_lbl.width, cell_val.width = Inches(2.2), Inches(4.3)
            set_cell_background(cell_lbl, HEX_LIGHT_BG)
            set_cell_margins(cell_lbl)
            set_cell_margins(cell_val)
            
            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_after = Pt(0)
            run_lbl = p_lbl.add_run(label)
            run_lbl.font.name = 'Calibri'
            run_lbl.bold = True
            run_lbl.font.color.rgb = COLOR_PRIMARY
            
            p_val = cell_val.paragraphs[0]
            p_val.paragraph_format.space_after = Pt(0)
            run_val = p_val.add_run(val)
            run_val.font.name = 'Calibri'
            run_val.font.color.rgb = COLOR_TEXT
            
        add_paragraph_with_spacing(doc, space_after=12)
        
        add_custom_heading(doc, s.get("benefitsTitle", "¿Cuáles son los beneficios?"), level=3, space_before=10, space_after=4)
        if s.get("benefitsIntro"):
            add_paragraph_with_spacing(doc, s["benefitsIntro"])
        for b in s["benefits"]:
            add_paragraph_with_spacing(doc, b, bullet=True, space_after=4)
            
        add_custom_heading(doc, s.get("recommendationsTitle", "Recomendaciones"), level=3, space_before=10, space_after=4)
        for rec in s["postOpRecommendations"]:
            add_paragraph_with_spacing(doc, rec, bullet=True, space_after=4)
            
        if s.get("risks") and len(s["risks"]) > 0:
            add_custom_heading(doc, s["risksTitle"], level=3, space_before=10, space_after=4)
            for risk in s["risks"]:
                add_paragraph_with_spacing(doc, risk, bullet=True, space_after=4)

        if s.get("seo"):
            add_custom_heading(doc, "Metadatos SEO", level=3, space_before=10, space_after=4)
            add_paragraph_with_spacing(doc, f"**Meta Title:** {s['seo'].get('title', '')}")
            add_paragraph_with_spacing(doc, f"**Meta Description:** {s['seo'].get('description', '')}")
            add_paragraph_with_spacing(doc, f"**Palabras Clave:** {', '.join(s['seo'].get('keywords', []))}")
            
        doc.add_paragraph()

    doc.add_page_break()

    # 1.4 Sección 3: Síntomas y Motivos de Consulta
    add_custom_heading(doc, "3. Síntomas y Motivos de Consulta", level=1, space_before=18, space_after=12)
    
    for idx, sym in enumerate(symptoms):
        add_custom_heading(doc, f"Síntoma {idx+1}: {sym['name']}", level=2, space_before=14, space_after=6)
        add_paragraph_with_spacing(doc, f"**Descripción:** {sym['description']}")
        
        add_custom_heading(doc, sym.get("causesTitle", "Causas comunes"), level=3, space_before=10, space_after=4)
        for c in sym["causes"]:
            add_paragraph_with_spacing(doc, c, bullet=True, space_after=4)
            
        add_custom_heading(doc, "¿Cuándo consultar al oftalmólogo?", level=3, space_before=10, space_after=4)
        add_paragraph_with_spacing(doc, sym["whyConsult"])
        
        if sym.get("seo"):
            add_custom_heading(doc, "Metadatos SEO", level=3, space_before=10, space_after=4)
            add_paragraph_with_spacing(doc, f"**Meta Title:** {sym['seo'].get('title', '')}")
            add_paragraph_with_spacing(doc, f"**Meta Description:** {sym['seo'].get('description', '')}")

    doc.add_page_break()

    # 1.5 Sección 4: Enfermedades y Padecimientos Tratados
    add_custom_heading(doc, "4. Enfermedades y Padecimientos Oculares", level=1, space_before=18, space_after=12)
    
    for idx, d in enumerate(diseases):
        add_custom_heading(doc, f"Padecimiento {idx+1}: {d['name']}", level=2, space_before=14, space_after=6)
        if d.get("technicalName"):
            add_paragraph_with_spacing(doc, f"**Nombre Médico/Técnico:** {d['technicalName']}")
        add_paragraph_with_spacing(doc, f"**Descripción:** {d['description']}")
        
        if d.get("mexicoStats"):
            add_callout(doc, d["mexicoStats"], title="Estadística e Impacto en México")
            
        add_custom_heading(doc, "Síntomas Principales", level=3, space_before=10, space_after=4)
        for s_item in d["symptoms"]:
            add_paragraph_with_spacing(doc, s_item, bullet=True, space_after=4)
            
        add_custom_heading(doc, "Causas y Origen", level=3, space_before=10, space_after=4)
        for c_item in d["causes"]:
            add_paragraph_with_spacing(doc, c_item, bullet=True, space_after=4)
            
        if d.get("riskFactors"):
            add_custom_heading(doc, "Factores de Riesgo", level=3, space_before=10, space_after=4)
            for rf in d["riskFactors"]:
                add_paragraph_with_spacing(doc, rf, bullet=True, space_after=4)
                
        if d.get("complications"):
            add_custom_heading(doc, "Posibles Complicaciones", level=3, space_before=10, space_after=4)
            for comp in d["complications"]:
                add_paragraph_with_spacing(doc, comp, bullet=True, space_after=4)
                
        add_custom_heading(doc, "Opciones de Tratamiento", level=3, space_before=10, space_after=4)
        for tr in d["treatments"]:
            add_paragraph_with_spacing(doc, tr, bullet=True, space_after=4)
            
        if d.get("faqs"):
            add_custom_heading(doc, "Preguntas Frecuentes (FAQs)", level=3, space_before=10, space_after=4)
            for faq in d["faqs"]:
                add_paragraph_with_spacing(doc, f"**P: {faq['question']}**")
                add_paragraph_with_spacing(doc, f"R: {faq['answer']}", space_after=8)
                
        if d.get("seo"):
            add_custom_heading(doc, "Metadatos SEO", level=3, space_before=10, space_after=4)
            add_paragraph_with_spacing(doc, f"**Meta Title:** {d['seo'].get('title', '')}")
            add_paragraph_with_spacing(doc, f"**Meta Description:** {d['seo'].get('description', '')}")

    doc.add_page_break()

    # 1.6 Sección 5: Aviso de Privacidad
    add_custom_heading(doc, "5. Aviso de Privacidad", level=1, space_before=18, space_after=12)
    add_paragraph_with_spacing(doc, privacy["intro"])
    for sec in privacy["sections"]:
        add_custom_heading(doc, sec["title"], level=2, space_before=10, space_after=4)
        add_paragraph_with_spacing(doc, sec["content"])

    # Guardar DOCX
    docx_filename = os.path.join(script_dir, "../../Dra_Fiona_Xacur_Contenido.docx")
    doc.save(docx_filename)
    print(f"[OK] Documento Word generado en: {docx_filename}")

    # ----------------------------------------------------
    # 2. GENERACIÓN DEL DOCUMENTO MARKDOWN (.md)
    # ----------------------------------------------------
    md_lines = []
    md_lines.append(f"# Guía de Contenidos Consolidados del Sitio Web")
    md_lines.append(f"## {doctor['title']} {doctor['name']}")
    md_lines.append(f"### {doctor['specialty']} | {doctor['subspecialty']}\n")
    md_lines.append(f"*Este documento consolida y organiza toda la información de textos, servicios, síntomas y enfermedades que componen el sitio web.*\n\n---\n")

    # 2.1 Perfil
    md_lines.append("## 1. Perfil Profesional de la Doctora\n")
    md_lines.append("### Biografía")
    md_lines.append(doctor['bio'] + "\n")
    md_lines.append("### Filosofía Médica")
    md_lines.append(f"> \"{doctor['philosophy']}\"\n")
    md_lines.append("### Información General de Consulta")
    md_lines.append(f"* **Dirección:** {doctor['address']}, {doctor['city']}, {doctor['state']}")
    md_lines.append(f"* **Teléfono:** {doctor['phone']}")
    md_lines.append(f"* **WhatsApp:** {doctor['whatsapp']}")
    md_lines.append(f"* **Email:** {doctor['email']}")
    md_lines.append(f"* **Precio de Consulta:** ${doctor['consultationPrice']} MXN")
    md_lines.append(f"* **Métodos de Pago:** {', '.join(doctor['paymentMethods'])}")
    md_lines.append(f"* **Aseguradoras:** {', '.join(doctor['insurances'])}")
    md_lines.append(f"* **Horarios:** {doctor['schedule']}")
    md_lines.append(f"* **Cédula Profesional:** {doctor['cedula']}")
    md_lines.append(f"* **Cédula de Especialidad:** {doctor['cedulaEspecialidad']}\n")

    md_lines.append("### Educación")
    for edu in doctor['education']:
        md_lines.append(f"* **{edu['degree']}** — {edu['institution']}" + (f" ({edu['year']})" if edu.get('year') else ""))
    md_lines.append("\n### Certificaciones")
    for cert in doctor['certifications']:
        md_lines.append(f"* **{cert['name']}** — {cert['institution']}")
    md_lines.append("\n### Experiencia Laboral")
    for exp in doctor['experience']:
        md_lines.append(f"* **{exp['year']} - {exp['title']}:** {exp['description']}")
    md_lines.append("\n---\n")

    # 2.2 Servicios
    md_lines.append("## 2. Portafolio de Servicios y Procedimientos Médicos\n")
    for idx, s in enumerate(services):
        md_lines.append(f"### Servicio {idx+1}: {s['name']}\n")
        md_lines.append(f"* **Descripción corta:** {s['description']}")
        md_lines.append(f"* **Descripción detallada:** {s['longDescription']}\n")
        md_lines.append("#### Ficha Técnica")
        md_lines.append(f"* **Tipo:** {s['type'].upper()}")
        md_lines.append(f"* **Anestesia:** {s['anesthesiaType']}")
        md_lines.append(f"* **Duración:** {s['duration']}")
        md_lines.append(f"* **Recuperación:** {s['recoveryTime']}")
        md_lines.append(f"* **¿Es doloroso?:** {'Sí' if s['isPainful'] else 'No'}")
        md_lines.append(f"* **Precio:** {s['priceRange']}")
        for k, v in s['technicalSpecs'].items():
            md_lines.append(f"* **{k}:** {v}")
        md_lines.append(f"\n#### {s.get('benefitsTitle', '¿Cuáles son los beneficios?')}")
        for b in s['benefits']:
            md_lines.append(f"* {b}")
        md_lines.append(f"\n#### {s.get('recommendationsTitle', 'Recomendaciones')}")
        for rec in s['postOpRecommendations']:
            md_lines.append(f"* {rec}")
        md_lines.append("\n---\n")

    # 2.3 Síntomas
    md_lines.append("## 3. Síntomas y Motivos de Consulta\n")
    for idx, sym in enumerate(symptoms):
        md_lines.append(f"### Síntoma {idx+1}: {sym['name']}\n")
        md_lines.append(f"**Descripción:** {sym['description']}\n")
        md_lines.append(f"#### {sym.get('causesTitle', 'Causas comunes')}")
        for c in sym['causes']:
            md_lines.append(f"* {c}")
        md_lines.append(f"\n**¿Cuándo consultar?:** {sym['whyConsult']}\n")
        md_lines.append("---\n")

    # 2.4 Enfermedades
    md_lines.append("## 4. Enfermedades y Padecimientos Oculares\n")
    for idx, d in enumerate(diseases):
        md_lines.append(f"### Padecimiento {idx+1}: {d['name']}")
        if d.get('technicalName'):
            md_lines.append(f"*Nombre Técnico: {d['technicalName']}*\n")
        md_lines.append(f"**Descripción:** {d['description']}\n")
        if d.get('mexicoStats'):
            md_lines.append(f"> **Estadística en México:** {d['mexicoStats']}\n")
        md_lines.append("#### Síntomas")
        for s_item in d['symptoms']:
            md_lines.append(f"* {s_item}")
        md_lines.append("\n#### Causas")
        for c_item in d['causes']:
            md_lines.append(f"* {c_item}")
        md_lines.append("\n#### Opciones de Tratamiento")
        for tr in d['treatments']:
            md_lines.append(f"* {tr}")
        if d.get('faqs'):
            md_lines.append("\n#### Preguntas Frecuentes")
            for faq in d['faqs']:
                md_lines.append(f"* **P: {faq['question']}**")
                md_lines.append(f"  *R: {faq['answer']}*")
        md_lines.append("\n---\n")

    # 2.5 Privacidad
    md_lines.append("## 5. Aviso de Privacidad\n")
    md_lines.append(privacy['intro'] + "\n")
    for sec in privacy['sections']:
        md_lines.append(f"### {sec['title']}")
        md_lines.append(sec['content'] + "\n")

    md_filename = os.path.join(script_dir, "../../Dra_Fiona_Xacur_Contenido.md")
    with open(md_filename, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_lines))
    print(f"[OK] Documento Markdown generado en: {md_filename}")

    # ----------------------------------------------------
    # 3. GENERACIÓN DEL DOCUMENTO HTML (.html)
    # ----------------------------------------------------
    html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Contenido Web - {doctor['title']} {doctor['name']}</title>
    <style>
        body {{ font-family: 'Segoe UI', system-ui, -apple-system, sans-serif; line-height: 1.6; color: #{HEX_TEXT}; background-color: #f8fafc; margin: 0; padding: 40px 20px; }}
        .container {{ max-width: 900px; margin: 0 auto; background: #fff; padding: 50px; border-radius: 16px; box-shadow: 0 4px 20px rgba(0,0,0,0.06); }}
        h1 {{ color: #{HEX_PRIMARY}; font-size: 28px; border-bottom: 3px solid #{HEX_ACCENT}; padding-bottom: 10px; }}
        h2 {{ color: #{HEX_PRIMARY}; font-size: 22px; margin-top: 30px; border-bottom: 1px solid #{HEX_BORDER}; padding-bottom: 6px; }}
        h3 {{ color: #{HEX_ACCENT}; font-size: 18px; margin-top: 20px; }}
        .quote {{ background: #{HEX_LIGHT_BG}; border-left: 4px solid #{HEX_ACCENT}; padding: 15px 20px; border-radius: 0 8px 8px 0; font-style: italic; color: #{HEX_PRIMARY}; font-weight: 500; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th, td {{ padding: 10px 14px; text-align: left; border-bottom: 1px solid #{HEX_BORDER}; }}
        th {{ background: #{HEX_LIGHT_BG}; color: #{HEX_PRIMARY}; font-weight: 600; }}
        ul {{ padding-left: 20px; }}
        li {{ margin-bottom: 6px; }}
        .badge {{ display: inline-block; padding: 4px 10px; background: #{HEX_LIGHT_BG}; color: #{HEX_PRIMARY}; border-radius: 20px; font-size: 12px; font-weight: bold; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Guía Completa de Contenido del Sitio Web</h1>
        <p><strong>Doctora:</strong> {doctor['title']} {doctor['name']}<br>
        <strong>Especialidad:</strong> {doctor['specialty']}<br>
        <strong>Subespecialidad:</strong> {doctor['subspecialty']}</p>
        
        <h2>1. Perfil Profesional</h2>
        <p>{doctor['bio']}</p>
        <div class="quote">"{doctor['philosophy']}"</div>
        
        <h3>Información del Consultorio</h3>
        <table>
            <tr><th>Campo</th><th>Información</th></tr>
            <tr><td>Dirección</td><td>{doctor['address']}</td></tr>
            <tr><td>Teléfono / WhatsApp</td><td>{doctor['phone']} / {doctor['whatsapp']}</td></tr>
            <tr><td>Email</td><td>{doctor['email']}</td></tr>
            <tr><td>Costo Consulta</td><td>${doctor['consultationPrice']} MXN</td></tr>
            <tr><td>Cédulas</td><td>Prof: {doctor['cedula']} | Esp: {doctor['cedulaEspecialidad']}</td></tr>
        </table>
        
        <h2>2. Portafolio de Servicios ({len(services)})</h2>
"""
    for s in services:
        html_content += f"""
        <h3>{s['name']}</h3>
        <p><strong>Descripción:</strong> {s['longDescription']}</p>
        <ul>
            <li><strong>Duración:</strong> {s['duration']} | <strong>Recuperación:</strong> {s['recoveryTime']}</li>
            <li><strong>Anestesia:</strong> {s['anesthesiaType']}</li>
        </ul>
        """
        
    html_content += f"""
        <h2>3. Síntomas ({len(symptoms)})</h2>
"""
    for sym in symptoms:
        html_content += f"""
        <h3>{sym['name']}</h3>
        <p>{sym['description']}</p>
        <p><em>¿Cuándo consultar?:</em> {sym['whyConsult']}</p>
        """

    html_content += f"""
        <h2>4. Enfermedades ({len(diseases)})</h2>
"""
    for d in diseases:
        html_content += f"""
        <h3>{d['name']}</h3>
        <p>{d['description']}</p>
        """

    html_content += """
    </div>
</body>
</html>
"""
    html_filename = os.path.join(script_dir, "../../Dra_Fiona_Xacur_Contenido.html")
    with open(html_filename, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"[OK] Documento HTML generado en: {html_filename}")

if __name__ == "__main__":
    main()
